"""
Document Generator Tool: Creates Word (.docx) files using python-docx.
"""

import logging
import re
from pathlib import Path

from docx import Document

from backend.tools.path_safety import safe_resolve_output_path

logger = logging.getLogger(__name__)

# Output directory for generated documents
OUTPUT_DIR = Path(__file__).parent.parent.parent / "workspace" / "outputs"


# Control characters that are invalid in XML (keep tab=0x09, LF=0x0A, CR=0x0D)
_XML_INVALID_RE = re.compile(r'[\x00-\x08\x0B\x0C\x0E-\x1F]')


def _sanitize_xml_str(text: str) -> str:
    """Remove XML-invalid control characters from a string."""
    return _XML_INVALID_RE.sub('', text)


def generate_doc(filename: str, title: str, content: str) -> str:
    """
    Generate a Word document with a title and body content.

    Args:
        filename: Name of the output file (e.g. "report.docx").
        title: The document title (rendered as Heading 1).
        content: The body text (rendered as a paragraph).

    Returns:
        Absolute path to the created file.
    """
    # Ensure output directory exists
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Resolve path with containment check
    output_path = safe_resolve_output_path(filename, OUTPUT_DIR)

    # Sanitize strings for XML compatibility
    safe_title = _sanitize_xml_str(str(title))
    safe_content = _sanitize_xml_str(str(content))

    try:
        doc = Document()

        # Add title as Heading 1
        doc.add_heading(safe_title, level=1)

        # Add content as a paragraph
        doc.add_paragraph(safe_content)

        # Save the document
        doc.save(str(output_path))

        logger.info(f"Document generated: {output_path}")
        return str(output_path)

    except Exception as e:
        logger.error(f"Document generation failed: {e}")
        raise
